"""populate_memo.py — Populate the Overland Posting Memo Word template.

Opens the bundled template, writes drafted content into the correct table cells
and paragraphs by editing existing runs in-place, and saves a populated copy.

Why in-place editing: the template uses numbered lists (Word auto-numbering) and
bold-prefix run patterns (e.g., bold "Revenue:" + plain "growth commentary") in
multi-paragraph cells. Clearing and re-adding paragraphs strips the numPr
properties and merges run-level formatting. In-place run replacement preserves
all of that.

A final `_cleanup_layout()` pass removes template cruft that becomes defects
once the template is populated: a duplicate empty paragraph that produces a
blank page between Recommendation and Designated Criteria, fixed minimum row
heights authored for the blank template, redundant header/footer references,
and tab characters in the running header (replaced with positional tabs for
robust right-alignment).

USAGE
-----
    python scripts/populate_memo.py <content.json> <output.docx> [--template <path>]

If --template is omitted, the script uses assets/posting-memo-template.docx
relative to the skill root.

CONTENT JSON SCHEMA
-------------------
All fields optional; omitted fields leave the template placeholder in place.
For multi-paragraph cells (strengths, considerations, discussion_analysis,
company_overview.bullets), the template ships with a fixed number of slots
matching expected content (5 for strengths/considerations, 6 for D&A and CO
bullets). Lists shorter than the slot count delete excess slots; lists longer
than the slot count clone the last slot's XML (preserving numbering style).

Bold-prefix bullets: when a paragraph has a bold "Header" run followed by a
plain ": detail" run, pass the full string ("Header: detail") and the script
will route the prefix into the bold run and the remainder into the plain run.

{
  "header": {
      "company_name": str, "owners": str, "hq_year": str,
      "sector_industry": str, "origination_source": str,
      "posting_team": str, "received_posted_date": str, "process_stage": str,
      "feedback_party": str, "feedback_deadline": str,
      "cover_company_name": str,    # short name for cover page + running header
                                    # (defaults to company_name minus parenthetical)
      "cover_date": str             # cover page date, e.g. 'March 15, 2026'
                                    # (defaults to received_posted_date)
  },
  "situation_overview": str,
  "company_overview": {"opening": str, "bullets": [str, ...]},
  "financial_headline": str,
  "discussion_analysis": [str, ...],
  "su_note": str,
  "risk_flags": {
      "conc": "Y"|"N"|"TBD", ..., "esg_rr": "n/a", "esg_sa": "n/a"
  },
  "strengths":      [str, ...],
  "considerations": [str, ...],
  "recommendation": str,
  "designated_criteria": {
      "ebitda": "Y"|"N"|"TBD", "leverage": ..., "secured": ...,
      "deal_size": ..., "yield": ..., "us_domiciled": ...,
      "other_considerations": str
  },
  "posting_rating": str,
  "final_rating": str | null
}
"""

from __future__ import annotations

import json
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph as _Paragraph

DEFAULT_TEMPLATE = (
    Path(__file__).resolve().parent.parent / "assets" / "posting-memo-template.docx"
)

RISK_ITEMS = [
    "conc", "cyclicality", "seasonality", "nwc_needs", "capex_needs",
    "project_based", "excessive_revolver", "bonding_surety",
    "fx_exposures", "raw_mat_volatility", "unique_accounting",
    "mgmt_issues", "regulatory_risks", "technology_risks",
]

RATING_ROW = {
    "very interesting": 2, "green": 2,
    "begin diligence": 3,  "yellow": 3,
    "high diligence bar": 4, "orange": 4,
    "no diligence path": 5, "red": 5,
    "alternative strategy": 6, "alt": 6, "alternative": 6,
}


# ---------------------------------------------------------------------------
# Run / paragraph helpers
# ---------------------------------------------------------------------------


def _remove_runs_after(paragraph: _Paragraph, keep: int) -> None:
    for run in list(paragraph.runs[keep:]):
        run._element.getparent().remove(run._element)


def _run_format_tuple(run) -> tuple:
    return (bool(run.bold), bool(run.italic), bool(run.underline))


def _find_format_boundary(paragraph: _Paragraph) -> int:
    """Index of the first run whose (bold, italic, underline) differs from
    run[0]'s. Returns len(runs) if all runs share the same formatting (i.e. no
    prefix/body split exists)."""
    runs = paragraph.runs
    if len(runs) < 2:
        return len(runs)
    base = _run_format_tuple(runs[0])
    for i in range(1, len(runs)):
        if _run_format_tuple(runs[i]) != base:
            return i
    return len(runs)


def _set_paragraph_text_smart(paragraph: _Paragraph, text: str) -> None:
    """Write text into a paragraph respecting existing run roles.

    Heuristic: if the paragraph has a header/body run pattern (run[0] uses some
    bold/italic/underline formatting that subsequent runs drop) AND the input
    contains ': ', split at the first ': '. The colon is included in the
    formatted prefix run so the visual styling matches the template's
    "Header:_body" convention. Multi-run prefixes (e.g., template runs 0-2 all
    italic+underline followed by run 3 plain) are collapsed into the first run.

    If formatting is uniform across runs, the entire text goes to run[0] and
    subsequent runs are removed.
    """
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(text)
        return

    boundary = _find_format_boundary(paragraph)
    # boundary is the first "body" run; runs[0:boundary] are the "header" runs.

    if 0 < boundary < len(runs) and ": " in text:
        prefix, _, body = text.partition(": ")
        # Include the colon in the formatted prefix run; leading space on body
        runs[0].text = prefix + ":"
        runs[boundary].text = " " + body
        # Remove header runs between 0 and the boundary (they were absorbed
        # into runs[0]) and any runs after the body run
        for r in list(runs[boundary + 1:]):
            r._element.getparent().remove(r._element)
        for r in list(runs[1:boundary]):
            r._element.getparent().remove(r._element)
    else:
        runs[0].text = text
        _remove_runs_after(paragraph, 1)


def _clone_paragraph_after(paragraph: _Paragraph) -> _Paragraph:
    """Deep-copy paragraph XML and insert immediately after."""
    new_el = deepcopy(paragraph._element)
    paragraph._element.addnext(new_el)
    return _Paragraph(new_el, paragraph._parent)


def _remove_paragraph(paragraph: _Paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)


# ---------------------------------------------------------------------------
# Cell writers
# ---------------------------------------------------------------------------


def set_cell_text(cell, text: str) -> None:
    """Single-line replacement preserving formatting of the first run."""
    if text is None:
        text = ""
    paragraphs = cell.paragraphs
    if not paragraphs:
        cell.add_paragraph(text)
        return
    for p in list(paragraphs[1:]):
        _remove_paragraph(p)
    _set_paragraph_text_smart(paragraphs[0], text)


def replace_paragraph_list(cell, items: list[str]) -> None:
    """Replace cell paragraphs with `items`, editing existing slots in place
    and preserving numbered-list properties. Excess slots are removed; if more
    items than slots, the last slot is cloned to preserve numbering style."""
    paragraphs = cell.paragraphs
    n_existing = len(paragraphs)
    n_new = len(items)

    if n_new == 0:
        for p in list(paragraphs[1:]):
            _remove_paragraph(p)
        if paragraphs:
            _set_paragraph_text_smart(paragraphs[0], "")
        return

    # Edit in place for slots that exist
    for i in range(min(n_new, n_existing)):
        _set_paragraph_text_smart(paragraphs[i], items[i])

    # Clone last paragraph if more items than slots
    if n_new > n_existing and n_existing > 0:
        anchor = paragraphs[n_existing - 1]
        for i in range(n_existing, n_new):
            anchor = _clone_paragraph_after(anchor)
            _set_paragraph_text_smart(anchor, items[i])

    # Remove excess slots if fewer items than slots
    elif n_new < n_existing:
        for p in list(paragraphs[n_new:]):
            _remove_paragraph(p)


def mark_x(cell) -> None:
    set_cell_text(cell, "X")


def clear_cell(cell) -> None:
    set_cell_text(cell, "")


def _populate_header_and_cover(doc, header: dict) -> None:
    """Populate the cover page (body[0] text frames) and the section running
    header. The template ships with literal placeholder strings ('[Company Name]',
    '[Month Day, Year]', '[Company]', '[Owner]') that we replace via per-run
    text substitution to preserve formatting.

    Cover page company name comes from 'cover_company_name'; falls back to
    'company_name' if absent. Cover date comes from 'cover_date'; falls back to
    'received_posted_date' if absent. The running header value '[Company] ([Owner])'
    is constructed as f"{cover_company_name} ({owners})" if both are provided.
    """
    cover_name = header.get("cover_company_name") or header.get("company_name") or ""
    cover_date = header.get("cover_date") or header.get("received_posted_date") or ""
    owners = header.get("owners", "")

    # Strip any trailing parenthetical from cover_name if it was reused from
    # company_name (which often looks like 'Stark Tech ("Stark", or the "Company")')
    if cover_name and "(" in cover_name and not header.get("cover_company_name"):
        cover_name = cover_name.split("(")[0].strip()

    # --- Cover page: replace literal placeholders in every <w:t> of body[0] ---
    if cover_name or cover_date:
        from docx.oxml.ns import qn
        body0 = doc.element.body[0]
        for t_el in body0.iter(qn("w:t")):
            if t_el.text == "[Company Name]" and cover_name:
                t_el.text = cover_name
            elif t_el.text == "[Month Day, Year]" and cover_date:
                t_el.text = cover_date

    # --- Section running header ---
    # The header placeholders ("[Company] ([Owner])", "[Month Day, Year]") are
    # split across many small runs because Word fragments runs at edit boundaries.
    # We work on the joined paragraph text, then consolidate post-tab runs into
    # a single run carrying the formatting of the first post-tab run.
    if not doc.sections:
        return
    hdr = doc.sections[0].header

    def _replace_after_tab(paragraph, replacement_after_tab: str) -> None:
        """If the paragraph has a tab, replace everything after the (last) tab
        with `replacement_after_tab`, consolidating into a single run that
        carries the formatting of the first post-tab run."""
        runs = paragraph.runs
        if not runs:
            return
        # Find the run containing the last tab
        tab_run_idx = None
        for i, r in enumerate(runs):
            if "\t" in r.text:
                tab_run_idx = i
        if tab_run_idx is None:
            return
        # Truncate the tab run to end at (and include) the tab
        tab_run = runs[tab_run_idx]
        tab_pos = tab_run.text.rfind("\t")
        tab_run.text = tab_run.text[: tab_pos + 1]
        # Take formatting from the run immediately after the tab run, if any,
        # otherwise from the tab run itself
        format_source = runs[tab_run_idx + 1] if tab_run_idx + 1 < len(runs) else tab_run
        # Remove all runs after the tab run
        for r in list(runs[tab_run_idx + 1 :]):
            r._element.getparent().remove(r._element)
        # Add a new run carrying the format source's formatting
        new_run = paragraph.add_run(replacement_after_tab)
        new_run.bold = format_source.bold
        new_run.italic = format_source.italic
        new_run.underline = format_source.underline
        if format_source.font.name:
            new_run.font.name = format_source.font.name
        if format_source.font.size:
            new_run.font.size = format_source.font.size

    # Para[0]: 'Overland Advantage\t[Company] ([Owner])'
    if cover_name and owners and len(hdr.paragraphs) >= 1:
        running_company = f"{cover_name} ({owners})"
        if "[Company]" in hdr.paragraphs[0].text:
            _replace_after_tab(hdr.paragraphs[0], running_company)
    # Para[1]: 'New Deal Posting Memo\t[Month Day, Year]'
    if cover_date and len(hdr.paragraphs) >= 2:
        if "[Month Day, Year]" in hdr.paragraphs[1].text:
            _replace_after_tab(hdr.paragraphs[1], cover_date)


# ---------------------------------------------------------------------------
# Layout cleanup (post-population)
# ---------------------------------------------------------------------------


def _has_visible_text(p_element) -> bool:
    """True if the paragraph element contains any non-whitespace text."""
    text = "".join(t.text or "" for t in p_element.iter(qn("w:t")))
    return bool(text.strip())


def _dedupe_consecutive_empty_paragraphs(body) -> None:
    """Remove the second (and subsequent) of any run of consecutive empty
    body-level paragraphs. The template has a duplicate empty paragraph between
    the Recommendation and Designated Criteria tables that causes a blank page;
    legitimate single empty paragraphs (used as gutters between section tables)
    are unaffected."""
    prev_was_empty_p = False
    to_remove = []
    for child in list(body):
        is_empty_p = (
            child.tag.endswith("}p") and not _has_visible_text(child)
        )
        if is_empty_p and prev_was_empty_p:
            to_remove.append(child)
            # Keep prev_was_empty_p True so a triple-run also gets cleaned
        else:
            prev_was_empty_p = is_empty_p
    for el in to_remove:
        body.remove(el)


def _set_row_min_height(row, twips: int) -> None:
    """Set a row's minimum height in twips (1440 = 1 inch). 20 effectively
    auto-sizes to content."""
    from docx.oxml import OxmlElement

    tr = row._tr
    trPr = tr.find(qn("w:trPr"))
    if trPr is None:
        trPr = OxmlElement("w:trPr")
        tr.insert(0, trPr)
    trHeight = trPr.find(qn("w:trHeight"))
    if trHeight is None:
        trHeight = OxmlElement("w:trHeight")
        trPr.append(trHeight)
    trHeight.set(qn("w:val"), str(twips))


def _consolidate_section_headers_footers(doc) -> None:
    """Remove even/first header & footer references from section properties,
    keeping only the default. The template's even/first headers/footers are
    empty; titlePg already suppresses the header/footer on the cover page."""
    body = doc.element.body
    sectPr = body.find(qn("w:sectPr"))
    if sectPr is None:
        for child in body:
            sp = child.find(qn("w:sectPr"))
            if sp is not None:
                sectPr = sp
                break
    if sectPr is None:
        return
    for ref_tag in ("headerReference", "footerReference"):
        for ref in list(sectPr.findall(qn(f"w:{ref_tag}"))):
            t = ref.get(qn("w:type"))
            if t in ("even", "first"):
                sectPr.remove(ref)


def _convert_header_tabs_to_ptab(doc) -> None:
    """Replace plain <w:tab/> elements in the section header with positional
    tabs anchored to the right margin. Robustly right-aligns the running header
    content regardless of left-side text length."""
    from docx.oxml import OxmlElement

    if not doc.sections:
        return
    hdr = doc.sections[0].header
    for paragraph in hdr.paragraphs:
        for r_el in paragraph._element.findall(qn("w:r")):
            for tab_el in list(r_el.findall(qn("w:tab"))):
                ptab = OxmlElement("w:ptab")
                ptab.set(qn("w:relativeTo"), "margin")
                ptab.set(qn("w:alignment"), "right")
                ptab.set(qn("w:leader"), "none")
                r_el.replace(tab_el, ptab)


def _cleanup_layout(doc) -> None:
    """Final-pass layout polish applied after all content has been written.

    The original template was authored with cosmetic visual sizing for the
    blank template (fixed row heights, duplicate empty paragraphs to space
    sections vertically, multiple header/footer references for first/even
    pages). Once the template is populated, those choices manifest as defects
    in the output: blank pages, excess whitespace below short content, and
    occasional layout drift. This pass corrects all of them.
    """
    body = doc.element.body

    # --- Finding 1: remove the duplicate empty paragraph that creates a blank
    # page between the Recommendation and Designated Criteria tables ---
    _dedupe_consecutive_empty_paragraphs(body)

    # --- Finding 2: reset fixed minimum row heights on populated content cells.
    # The template authored 864/2448/3600-twip minimums on populated cells for
    # visual sizing of the blank template; once cells are populated those
    # minimums create excess whitespace. Reset to 20 (effectively auto-size).
    # The Strengths/Considerations row is left tall (so the Recommendation
    # table sits at the page bottom) but reduced from 7488 to 7344 — that 144
    # twip headroom is what allows the layout to flow without forcing an extra
    # blank page after the Recommendation table's hard page break. ---
    HEIGHT_RESETS = [
        (1, 1, 20),    # Situation Overview cell
        (2, 1, 20),    # Company Overview merged cell
        (2, 2, 20),    # LTM Headline row
        (2, 4, 20),    # Discussion & Analysis cell
        (5, 1, 7344),  # Strengths/Considerations row (kept tall, slightly relaxed)
    ]
    for ti, ri, twips in HEIGHT_RESETS:
        if ti < len(doc.tables) and ri < len(doc.tables[ti].rows):
            _set_row_min_height(doc.tables[ti].rows[ri], twips)

    # --- Finding 3: simplify section properties to a single default header
    # and footer reference (the template's even/first variants are empty) ---
    _consolidate_section_headers_footers(doc)

    # --- Finding 4: replace regular tabs in the header with positional tabs
    # anchored to the right margin, ensuring robust right-alignment of the
    # company name and date regardless of left-side text length ---
    _convert_header_tabs_to_ptab(doc)


# ---------------------------------------------------------------------------
# Population
# ---------------------------------------------------------------------------


def populate(template_path: Path, content: dict, output_path: Path) -> None:
    doc = Document(str(template_path))

    # --- Cover page + running header (body[0] text frames + section header) ---
    h = content.get("header", {})
    _populate_header_and_cover(doc, h)

    # --- Section 1: Deal header (Table 0) ---
    t = doc.tables[0]
    pairs = [
        (0, 1, "company_name"),       (0, 3, "posting_team"),
        (1, 1, "owners"),              (1, 3, "received_posted_date"),
        (2, 1, "hq_year"),             (2, 3, "process_stage"),
        (3, 1, "sector_industry"),     (3, 3, "feedback_party"),
        (4, 1, "origination_source"), (4, 3, "feedback_deadline"),
    ]
    for row, col, key in pairs:
        if key in h:
            set_cell_text(t.rows[row].cells[col], h[key])

    # --- Section 2: Situation Overview (Table 1) ---
    if "situation_overview" in content:
        set_cell_text(doc.tables[1].rows[1].cells[0], content["situation_overview"])

    # --- Sections 3–5: Company Overview + Headline + D&A (Table 2) ---
    t = doc.tables[2]
    co = content.get("company_overview")
    if co:
        # Cell [1,0] (merged): para[0] = opening (plain), paras[1..6] = bullets (numbered)
        co_cell = t.rows[1].cells[0]
        opening = co.get("opening", "")
        bullets = list(co.get("bullets", []))
        # The template has 1 opening paragraph + 6 numbered bullet slots = 7 paragraphs total.
        # We keep that structure: opening goes to paragraphs[0], bullets go to paragraphs[1:].
        existing = co_cell.paragraphs
        if existing:
            # Update opening in para[0]
            if opening:
                _set_paragraph_text_smart(existing[0], opening)
            # Update bullets starting at para[1]
            n_bullet_slots = len(existing) - 1
            n_bullets = len(bullets)
            for i in range(min(n_bullets, n_bullet_slots)):
                _set_paragraph_text_smart(existing[1 + i], bullets[i])
            if n_bullets > n_bullet_slots and n_bullet_slots > 0:
                anchor = existing[-1]
                for i in range(n_bullet_slots, n_bullets):
                    anchor = _clone_paragraph_after(anchor)
                    _set_paragraph_text_smart(anchor, bullets[i])
            elif n_bullets < n_bullet_slots:
                # Remove excess bullet paragraphs
                paragraphs_after_update = co_cell.paragraphs
                for p in list(paragraphs_after_update[1 + n_bullets:]):
                    _remove_paragraph(p)
    if "financial_headline" in content:
        set_cell_text(t.rows[2].cells[0], content["financial_headline"])
    if "discussion_analysis" in content:
        replace_paragraph_list(t.rows[4].cells[1], content["discussion_analysis"])

    # --- Section 6: S&U Note (first 'Note: [X]' paragraph in body) ---
    su_note = content.get("su_note")
    if su_note:
        for p in doc.paragraphs:
            if p.text.strip() == "Note: [X]":
                _set_paragraph_text_smart(p, su_note)
                break

    # --- Section 7: Risk Flags (Table 4) ---
    t = doc.tables[4]
    flags = content.get("risk_flags", {})
    for idx, item in enumerate(RISK_ITEMS, start=1):
        flag = (flags.get(item) or "TBD").upper()
        for r in (1, 2, 3):
            clear_cell(t.rows[r].cells[idx])
        target_row = {"Y": 1, "N": 2, "TBD": 3}.get(flag, 3)
        mark_x(t.rows[target_row].cells[idx])
    if "esg_rr" in flags or "esg_sa" in flags:
        rr = flags.get("esg_rr", "n/a")
        sa = flags.get("esg_sa", "n/a")
        # No '|' separator: the cell is narrow (~0.4") so the text always wraps
        # to two lines. Inserting '|' creates an orphan at the start of line 2.
        set_cell_text(t.rows[1].cells[15], f"RR: {rr} SA: {sa}")

    # --- Sections 8 & 9: Strengths / Considerations (Table 5) ---
    t = doc.tables[5]
    if "strengths" in content:
        replace_paragraph_list(t.rows[1].cells[0], content["strengths"])
    if "considerations" in content:
        replace_paragraph_list(t.rows[1].cells[1], content["considerations"])

    # --- Section 10: Recommendation (Table 6) ---
    if "recommendation" in content:
        set_cell_text(doc.tables[6].rows[1].cells[0], content["recommendation"])

    # --- Section 11: Designated Criteria + Posting Rating (Table 7) ---
    t = doc.tables[7]
    crit = content.get("designated_criteria", {})
    for key, row in [("ebitda", 2), ("leverage", 3), ("secured", 4)]:
        flag = (crit.get(key) or "TBD").upper()
        for col in (1, 2, 3):
            clear_cell(t.rows[row].cells[col])
        col_for = {"Y": 1, "N": 2, "TBD": 3}.get(flag, 3)
        mark_x(t.rows[row].cells[col_for])
    for key, row in [("deal_size", 2), ("yield", 3), ("us_domiciled", 4)]:
        flag = (crit.get(key) or "TBD").upper()
        for col in (5, 6, 7):
            clear_cell(t.rows[row].cells[col])
        col_for = {"Y": 5, "N": 6, "TBD": 7}.get(flag, 7)
        mark_x(t.rows[row].cells[col_for])
    if "other_considerations" in crit:
        set_cell_text(t.rows[6].cells[0], crit["other_considerations"])
    rating = (content.get("posting_rating") or "").lower().strip()
    if rating:
        for row in (2, 3, 4, 5, 6):
            clear_cell(t.rows[row].cells[10])
        if rating in RATING_ROW:
            mark_x(t.rows[RATING_ROW[rating]].cells[10])

    # --- Section 12: Final Rating (Table 8) — leave default for pre-IC ---
    if content.get("final_rating"):
        set_cell_text(doc.tables[8].rows[1].cells[0], content["final_rating"])

    # --- Final pass: layout cleanup ---
    # Removes template cruft (duplicate empty paragraph causing a blank page,
    # fixed minimum row heights designed for the blank template, redundant
    # even/first header references) and applies header tab polish.
    _cleanup_layout(doc)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------


def _parse_args(argv: list[str]) -> tuple[Path, Path, Path]:
    if len(argv) < 3:
        sys.stderr.write(
            "usage: populate_memo.py <content.json> <output.docx> [--template <path>]\n"
        )
        sys.exit(2)
    content_path = Path(argv[1])
    output_path = Path(argv[2])
    template_path = DEFAULT_TEMPLATE
    if "--template" in argv:
        i = argv.index("--template")
        template_path = Path(argv[i + 1])
    return content_path, output_path, template_path


def main(argv: list[str]) -> None:
    content_path, output_path, template_path = _parse_args(argv)
    if not template_path.exists():
        sys.stderr.write(f"ERROR: template not found at {template_path}\n")
        sys.exit(2)
    if not content_path.exists():
        sys.stderr.write(f"ERROR: content JSON not found at {content_path}\n")
        sys.exit(2)
    with open(content_path, "r") as f:
        content = json.load(f)
    populate(template_path, content, output_path)
    print(f"Wrote populated memo: {output_path}")


if __name__ == "__main__":
    main(sys.argv)
