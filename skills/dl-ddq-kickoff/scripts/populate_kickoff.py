"""populate_kickoff.py - Populate the Wells & Overland Kick-Off Data Requests
one-pager Word template in place.

The template is a flat, single-page bulleted list (no tables): three numbered
section headers (Historical Financials, Historical KPIs, Forecasted & Budgeted
Financials) each followed by `Bullet2` sub-bullets carrying bracket
placeholders (FY'[YY], LTM [M/YY], Q[X]'[YY]). Two sub-bullets carry footnote
references that must survive editing.

Why in-place editing: the bullets use the `Bullet2` paragraph style and a
`SubtleReference` run style, and two carry `<w:footnoteReference>` runs.
Regenerating the document would strip the list styling and drop the footnotes.
This script edits existing runs and clones existing list paragraphs so every
style, the numbering, and both footnotes are preserved.

USAGE
-----
    python scripts/populate_kickoff.py <content.json> <output.docx> \
        [--template <path>]

If --template is omitted the bundled
assets/kickoff-data-requests-template.docx is used.

CONTENT JSON SCHEMA
-------------------
{
  "company_name": str,                 # header company placeholder
  "owner": str | null,                 # header "(Sponsor)" parenthetical;
                                        #   null -> the parenthetical is removed
  "as_of_date": str | null,            # e.g. "May 28, 2026"; refreshes the
                                        #   header DATE field's cached value
  "periods": {                          # from scripts/compute_periods.py
      "audited_range":   str,           #   "FY'23-'25"
      "ltm_anchor":      str,           #   "3/26"
      "quarterly_range": str,           #   "FY'24 - Q1'26"
      "budget_fy":       str,           #   "FY'26"
      "forecast_range":  str            #   "FY'26-'30"
  },
  "compliance_cert_applicable": bool,  # False -> the existing compliance-cert
                                        #   line is suffixed N/A, never deleted
  "stock_cut_requests":     [str, ...], # every-borrower analytical cuts
  "borrower_kpi_requests":  [str, ...]  # sector-specific KPI lines (capped)
}

Boilerplate text around each period fragment is owned by THIS script (single
source of truth); compute_periods.py owns the dates. Keeping the two split
means the period math and the request wording each have exactly one home.
"""

from __future__ import annotations

import json
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph as _Paragraph

DEFAULT_TEMPLATE = (
    Path(__file__).resolve().parent.parent
    / "assets"
    / "kickoff-data-requests-template.docx"
)

# One-page guard. The template body + numbered headers + the standard set fit a
# single page at the template's styling with ~14 terse noun-phrase KPI lines.
# Past that the list spills to a second page, which defeats the one-pager.
# Tighten phrasing rather than raising this; it is a layout fact, not a knob.
MAX_KPI_LINES = 14


# ---------------------------------------------------------------------------
# Run / paragraph helpers
# ---------------------------------------------------------------------------


def _run_has_footnote(run) -> bool:
    return run._element.find(qn("w:footnoteReference")) is not None


def _iter_body_paragraphs(doc):
    return doc.paragraphs


def _set_line_text_preserving_footnote(paragraph: _Paragraph, text: str) -> None:
    """Write `text` into the paragraph's first text run, drop the other text
    runs, and leave any footnote-reference run in place at the end.

    Every run in these bullets shares the same `SubtleReference` run style, so
    collapsing the text into run[0] preserves the visible formatting. A
    footnote run carries no `w:t` and is never touched, so footnotes 1 and 2
    survive. A single trailing space is added before a surviving footnote so
    the superscript mark does not abut the text.
    """
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(text)
        return

    footnote_runs = [r for r in runs if _run_has_footnote(r)]
    text_runs = [r for r in runs if not _run_has_footnote(r)]

    if text_runs:
        carrier = text_runs[0]
        carrier.text = text + (" " if footnote_runs else "")
        for r in text_runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        # No text run at all (degenerate); prepend one before the footnote.
        paragraph.add_run(text + " ")
    # footnote_runs are left exactly where they were (after the carrier).


def _clone_bullet_after(anchor: _Paragraph) -> _Paragraph:
    """Deep-copy the anchor Bullet2 paragraph, strip footnote and proofErr
    cruft from the copy, and insert it immediately after the anchor. Returns
    the new paragraph (text not yet set)."""
    new_el = deepcopy(anchor._element)
    # Remove any footnote-reference runs cloned from the anchor.
    for r_el in list(new_el.findall(qn("w:r"))):
        if r_el.find(qn("w:footnoteReference")) is not None:
            new_el.remove(r_el)
    # Drop proofErr markers (they reference the original text spans).
    for tag in ("w:proofErr",):
        for el in list(new_el.findall(qn(tag))):
            new_el.remove(el)
    anchor._element.addnext(new_el)
    return _Paragraph(new_el, anchor._parent)


def _find_paragraph(doc, needle: str) -> _Paragraph:
    """Return the first body paragraph whose joined text contains `needle`.
    Raises if absent so a template change fails loudly instead of silently
    skipping a required line."""
    for p in _iter_body_paragraphs(doc):
        if needle in p.text:
            return p
    raise ValueError(
        f"template anchor not found: {needle!r}. The bundled template may "
        f"have changed; populate_kickoff.py must be re-checked against it."
    )


# ---------------------------------------------------------------------------
# Header
# ---------------------------------------------------------------------------


def _populate_header(doc, company_name: str, owner: str | None, as_of: str | None) -> None:
    """Replace the header placeholders. The default header (header2) reads
    'Wells Fargo & Overland Advantage  <tab>  Company Name (Sponsor)' on line 1
    and 'Kick-Off Data Request List <tab> <DATE field>' on line 2.

    'Company Name' and 'Sponsor' are discrete runs we substitute by text. When
    `owner` is None the '(Sponsor)' parenthetical (the ' (', 'Sponsor', ')'
    runs) is removed. The DATE field auto-updates on open; we also refresh its
    cached display value so the document reads correctly before recalculation.
    """
    if not doc.sections:
        return
    hdr = doc.sections[0].header
    for p in hdr.paragraphs:
        runs = p.runs
        for i, r in enumerate(runs):
            if r.text == "Company Name" and company_name:
                r.text = company_name
            elif r.text == "Sponsor":
                if owner:
                    r.text = owner
                else:
                    # Remove ' (' , 'Sponsor', ')' to drop the parenthetical.
                    for j in (i - 1, i, i + 1):
                        if 0 <= j < len(runs) and runs[j].text in (
                            " (",
                            "Sponsor",
                            ")",
                        ):
                            runs[j].text = ""
        # Refresh the cached DATE field display value. The field auto-updates
        # on open; refreshing the cached run keeps the document correct before
        # recalculation. The cached value is the run between the 'separate' and
        # 'end' fldChar and reads like 'Month D, YYYY'.
        if as_of:
            for r in p.runs:
                if r.text and _looks_like_cached_date(r.text):
                    r.text = as_of


def _looks_like_cached_date(text: str) -> bool:
    """Heuristic: the DATE field's cached value looks like 'Month D, YYYY'."""
    parts = text.replace(",", "").split()
    return len(parts) == 3 and parts[0].isalpha() and parts[2].isdigit()


# ---------------------------------------------------------------------------
# Line builders (this script owns the boilerplate; periods come from JSON)
# ---------------------------------------------------------------------------


def _build_lines(periods: dict, compliance_applicable: bool) -> dict[str, str]:
    audited = periods["audited_range"]
    ltm = periods["ltm_anchor"]
    qrange = periods["quarterly_range"]
    budget = periods["budget_fy"]
    forecast = periods["forecast_range"]

    compliance_base = (
        f"{qrange} Quarterly Existing Loan Reporting & Compliance Certificates"
    )
    compliance_line = compliance_base + (
        " (If Applicable)"
        if compliance_applicable
        else " — N/A (no existing reporting facility)"
    )

    return {
        "Audited Financial Statements": f"{audited} Audited Financial Statements",
        "Bridge to Consolidated EBITDA": (
            f"LTM {ltm} Income Statement & Bridge to Consolidated EBITDA"
        ),
        "Quarterly Internal Financial": (
            f"{qrange} Quarterly Internal Financial Statements "
            f"(i.e., Income Statement, Balance Sheet, Cash Flow Statement)"
        ),
        "Loan Reporting": compliance_line,
        "Quarterly Mgmt. KPI": f"{qrange} Quarterly Mgmt. KPI’s",
        "Budgeted Financial Statements": f"{budget} Budgeted Financial Statements",
        "Long-Term Financial Statement Forecast": (
            f"{forecast} Long-Term Financial Statement Forecast"
        ),
    }


# ---------------------------------------------------------------------------
# Population
# ---------------------------------------------------------------------------


def populate(template_path: Path, content: dict, output_path: Path) -> None:
    doc = Document(str(template_path))

    company = content.get("company_name", "")
    owner = content.get("owner")
    as_of = content.get("as_of_date")
    _populate_header(doc, company, owner, as_of)

    periods = content.get("periods")
    if not periods:
        raise ValueError(
            "content['periods'] is required (run scripts/compute_periods.py "
            "and pass its output as content['periods'])."
        )
    compliance_applicable = bool(content.get("compliance_cert_applicable", False))
    lines = _build_lines(periods, compliance_applicable)

    # Rewrite the seven standard period lines in place.
    for anchor, new_text in lines.items():
        para = _find_paragraph(doc, anchor)
        _set_line_text_preserving_footnote(para, new_text)

    # Insert the standard "stock-cut" lines and the borrower-specific KPI
    # lines as new Bullet2 sub-bullets under the Historical KPIs section,
    # after the (now rewritten) Quarterly Mgmt. KPI's line.
    stock_cuts = list(content.get("stock_cut_requests", []))
    kpi_lines = list(content.get("borrower_kpi_requests", []))

    if len(kpi_lines) > MAX_KPI_LINES:
        raise ValueError(
            f"borrower_kpi_requests has {len(kpi_lines)} lines; the one-page "
            f"ceiling is {MAX_KPI_LINES}. Tighten the KPI phrasing into terse "
            f"noun phrases and merge related cuts rather than raising the cap. "
            f"Overflow lines:\n  - " + "\n  - ".join(kpi_lines[MAX_KPI_LINES:])
        )

    kpi_anchor = _find_paragraph(doc, "Quarterly Mgmt. KPI")
    cursor = kpi_anchor
    for line in stock_cuts + kpi_lines:
        cursor = _clone_bullet_after(cursor)
        _set_line_text_preserving_footnote(cursor, line)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------


def _parse_args(argv: list[str]) -> tuple[Path, Path, Path]:
    if len(argv) < 3:
        sys.stderr.write(
            "usage: populate_kickoff.py <content.json> <output.docx> "
            "[--template <path>]\n"
        )
        sys.exit(2)
    content_path = Path(argv[1])
    output_path = Path(argv[2])
    template_path = DEFAULT_TEMPLATE
    if "--template" in argv:
        template_path = Path(argv[argv.index("--template") + 1])
    return content_path, output_path, template_path


def main(argv: list[str]) -> None:
    content_path, output_path, template_path = _parse_args(argv)
    if not template_path.exists():
        sys.stderr.write(f"ERROR: template not found at {template_path}\n")
        sys.exit(2)
    if not content_path.exists():
        sys.stderr.write(f"ERROR: content JSON not found at {content_path}\n")
        sys.exit(2)
    with open(content_path, "r", encoding="utf-8") as f:
        content = json.load(f)
    try:
        populate(template_path, content, output_path)
    except ValueError as exc:
        sys.stderr.write(f"ERROR: {exc}\n")
        sys.exit(2)
    print(f"Wrote populated kick-off data request: {output_path}")


if __name__ == "__main__":
    main(sys.argv)
